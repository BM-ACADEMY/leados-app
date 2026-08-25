from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("Leados-AllianceOS-Subscription-Onboarding-Plan-and-Gap-Analysis.docx")


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text, bold=False, color=None, size=8.5):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    repeat_header(t.rows[0])
    for i, h in enumerate(headers):
        set_cell_text(t.rows[0].cells[i], h, True, "FFFFFF", 8.5)
        shade(t.rows[0].cells[i], "1F4E78")
    for ridx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if ridx % 2:
                shade(cells[i], "EAF2F8")
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return t


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.65)
sec.bottom_margin = Inches(0.65)
sec.left_margin = Inches(0.65)
sec.right_margin = Inches(0.65)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(9.5)
styles["Title"].font.name = "Aptos Display"
styles["Title"].font.size = Pt(25)
styles["Title"].font.color.rgb = RGBColor(31, 78, 120)
for s in ("Heading 1", "Heading 2", "Heading 3"):
    styles[s].font.name = "Aptos Display"
    styles[s].font.color.rgb = RGBColor(31, 78, 120)

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("LeadOS + AllianceOS\nCRM Subscription Product Plan")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Onboarding options, commercial choices, current-state audit, missing-process checklist, and implementation blueprint")
r.bold = True
r.font.size = Pt(12)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run(f"Decision draft | {date.today().strftime('%d %B %Y')} | Scope: LeadOS and AllianceOS only")

box = doc.add_table(rows=1, cols=1)
box.style = "Table Grid"
shade(box.cell(0, 0), "FFF2CC")
set_cell_text(box.cell(0, 0), "Important: prices in this document are illustrative decision options, not approved production prices. Mafiya OS, Thedal OS, and Content OS are excluded from this product scope. Existing files for those modules were reviewed only where needed to distinguish reusable patterns; they are not proposed as customer entitlements.", True, "7F6000", 9.5)

doc.add_heading("1. Executive decision summary", level=1)
doc.add_paragraph("The present portal is an internal, authenticated operating application—not yet a complete multi-tenant SaaS product. LeadOS and AllianceOS functionality exists, but the commercial customer journey around it is incomplete. The recommended launch model is a hybrid onboarding process: customers can choose a package and submit registration/payment online, while the operations team verifies integrations and releases credentials. This reduces onboarding delay without automatically exposing sensitive provider credentials.")
table(doc, ["Decision", "Recommended launch choice", "Alternatives retained"], [
    ("Product scope", "LeadOS, AllianceOS, and a combined Growth Suite", "Sell only the bundle; or quote every client individually"),
    ("Term", "1 month, 3 months, 1 year; custom enterprise term", "Add 2-year public term after retention data exists"),
    ("Payment", "Razorpay online plus controlled offline recording", "Online only; or fully manual invoice collection"),
    ("Onboarding", "Hybrid assisted automation", "Self-service; team-managed form/file; fully manual"),
    ("Access", "Tenant-isolated account with module entitlements", "Separate deployment per customer (higher cost)"),
    ("Go-live rule", "Payment verified + compliance/integration checklist complete", "Immediate access with restricted setup mode"),
])

doc.add_heading("2. Product packaging options", level=1)
doc.add_paragraph("Use module entitlements so the same portal can safely show only what the customer purchased. Do not remove unrelated source modules merely to create a package; hide and block them through server-side entitlements, navigation rules, and authorization checks.")
table(doc, ["Option", "Who it fits", "Included", "Commercial method"], [
    ("A — LeadOS", "Teams managing inbound leads, WhatsApp follow-up, campaigns and conversion", "Lead capture, inbox, pipeline/leads, campaigns, templates, AI knowledge, reports; agreed integrations", "Fixed plan tiers + usage limits"),
    ("B — AllianceOS", "Partnership/outbound teams managing prospecting and outreach", "Prospect upload/list, outreach planner, email/WhatsApp campaigns, replies, knowledge/prompts, analytics", "Fixed plan tiers + sender/contact limits"),
    ("C — Growth Suite (recommended bundle)", "Customers needing inbound CRM plus alliance/outbound workflow", "LeadOS + AllianceOS with shared account and reporting boundary", "Bundle discount compared with two separate subscriptions"),
    ("D — Custom Enterprise", "High volume, custom integrations, data migration, SLA or dedicated environment", "Negotiated modules, limits, onboarding, support and security terms", "Written quotation/order form"),
])

doc.add_heading("3. Illustrative pricing worksheet (INR)", level=1)
doc.add_paragraph("These numbers are starting hypotheses for approval. Validate cloud, WhatsApp/Meta, email, AI, data-provider, onboarding, support and GST costs before publishing. Provider usage charges should be stated as included up to a limit or billed separately.")
table(doc, ["Package", "Starter / month", "Growth / month", "Scale / month", "Suggested control"], [
    ("LeadOS", "₹7,999", "₹14,999", "₹29,999", "Leads, users, messages, campaigns, AI usage"),
    ("AllianceOS", "₹9,999", "₹17,999", "₹34,999", "Prospects, senders, emails/WhatsApp, campaigns, AI usage"),
    ("Growth Suite", "₹14,999", "₹27,999", "₹54,999", "Combined limits; 10–20% bundle advantage"),
    ("Enterprise", "Custom", "Custom", "Custom", "Minimum commitment + SLA + overages"),
])
table(doc, ["Term", "Pricing formula", "Example on ₹14,999 monthly", "Renewal behavior"], [
    ("1 month", "Monthly price × 1", "₹14,999", "Auto-renew online or manual renewal"),
    ("3 months", "Monthly price × 3 less 5%", "₹42,747 (rounded)", "Renew every 3 months"),
    ("1 year", "Monthly price × 12 less 15%", "₹152,990 (rounded)", "Annual advance; recommended value plan"),
    ("2 years", "Monthly price × 24 less 20%", "₹287,981 (rounded)", "Use only with price-rise, SLA and exit clauses"),
    ("Custom", "Selected dates / seats / usage / services", "Quotation", "Manual approval; optional milestone billing"),
])
bullet(doc, "Optional one-time setup fee: ₹5,000–₹25,000 based on migration, Meta/email setup, training and workflow customization.")
bullet(doc, "Choose whether prices are GST-inclusive or GST-exclusive and show this consistently on pricing, checkout, invoice and renewal notices.")
bullet(doc, "Do not promise unlimited WhatsApp, email, AI or storage usage unless a written fair-use and overage policy exists.")

doc.add_heading("4. Payment modes and controls", level=1)
table(doc, ["Mode", "Customer experience", "Back-office control", "Recommended status flow"], [
    ("Online one-time", "Checkout/payment link via Razorpay: UPI, card, net banking and supported methods", "Verify signature/webhook; create payment and invoice record", "pending → paid → provisioning → active"),
    ("Online recurring", "Razorpay subscription/mandate with consent", "Store provider customer/subscription IDs; process renewal/failure webhooks", "trialing/active → past_due → suspended/cancelled"),
    ("Offline bank/UPI", "Customer receives quotation/invoice and pays by transfer", "Staff records reference, date, amount and proof; finance approves", "awaiting_verification → paid → provisioning → active"),
    ("Cheque/cash", "Available only if business policy permits", "Receipt number, collector, deposit/clearance and approval audit", "received → cleared → paid; never activate before clearance"),
    ("Custom milestones", "Deposit plus scheduled installments", "Order form defines entitlement and suspension rules", "partially_paid → active/restricted according to contract"),
])
doc.add_paragraph("Required safeguards: idempotent webhooks, signature verification, unique provider transaction ID, immutable audit log, reconciliation report, refund/credit-note workflow, invoice numbering, failed-payment retry, grace period, and maker-checker approval for offline entries.")

doc.add_heading("5. Client onboarding models", level=1)
table(doc, ["Model", "Flow", "Advantages", "Risks / controls"], [
    ("A — Self-service", "Register → verify email/phone → choose plan → accept terms → pay → create tenant/user → setup wizard", "Fast and scalable", "Requires mature tenant isolation, fraud controls, automated entitlements and support content"),
    ("B — Assisted hybrid (recommended)", "Customer registers/chooses/pays → system creates provisioning case → team verifies integrations → customer sets password → go-live", "Balances automation with Meta/email setup realities", "Define SLA, owner, checklist and reminders so cases do not stall"),
    ("C — Team-created from form", "Customer submits secure form → sales/ops validates → finance records payment → admin creates tenant and sends invite", "Good for early-stage launch and custom contracts", "Manual errors; requires dual checks and audit trail"),
    ("D — Bulk file intake", "Authorized team uploads validated CSV/XLSX template → preview/errors → approve → create accounts → send invites", "Useful for partners or many client users", "Never accept passwords/secrets in file; provide row-level error report and rollback"),
    ("E — Dedicated deployment", "Separate environment/database and managed setup", "Strong isolation and customization", "Highest infrastructure, release and support cost"),
])

doc.add_heading("6. Recommended end-to-end onboarding workflow", level=1)
for text_item in [
    "Customer opens public product/pricing page and selects LeadOS, AllianceOS, Growth Suite, or Request Custom Quote.",
    "Customer registers owner name, business/legal name, email, mobile, country, GST details (if applicable), expected usage and consent. Email/phone is verified.",
    "System creates a prospect/application record—not an active tenant—and captures terms/version acceptance.",
    "Customer chooses duration and payment mode. Online payment uses a provider checkout; offline payment creates an invoice and verification task.",
    "Verified payment creates an order, payment allocation, subscription with exact start/end dates, and purchased module entitlements in one transactional workflow.",
    "System creates the tenant and owner invitation. Send a one-time, expiring password-setup link; never email a reusable plaintext password.",
    "Provisioning checklist collects integrations: WhatsApp/Meta, email sender, domain/DNS, branding, imports, AI knowledge and user invitations. Secrets go into a secret manager/encrypted store, not form files or email.",
    "Operations verifies provider connections and completes a test lead/message/campaign using a safe test target.",
    "Customer receives welcome email, invoice/receipt, login URL, setup status, support contacts and training link. Account changes from provisioning to active.",
    "Automations send expiry notices (30/15/7/1 days), retry failed renewals, apply the approved grace period, restrict paid actions after expiry, and preserve read/export access according to policy.",
]:
    numbered(doc, text_item)

doc.add_heading("7. Registration form and file options", level=1)
doc.add_heading("Minimum registration/form fields", level=2)
table(doc, ["Section", "Required fields", "Notes"], [
    ("Owner", "Name, work email, mobile, designation", "Verify email; verify mobile when used for alerts"),
    ("Business", "Legal/display name, address, country/state, website/domain, industry", "GSTIN/PAN only when commercially required; validate format"),
    ("Product", "Module, tier, duration, requested start date, expected users/usage", "Custom selection routes to quotation approval"),
    ("Billing", "Billing name/address/email, GST treatment, payment mode", "Do not collect card/bank credentials directly"),
    ("Integrations", "Meta business/WABA readiness, WhatsApp number, email/domain readiness", "Collect IDs through wizard; store tokens securely"),
    ("Legal", "Terms, privacy, DPA/consent, marketing consent separately", "Store document version, timestamp, IP and actor"),
])
doc.add_heading("Bulk upload template", level=2)
doc.add_paragraph("Recommended columns: external_reference, business_name, owner_name, owner_email, owner_phone, country, package_code, plan_code, term_code, start_date, billing_email, gstin, payment_mode, payment_reference, amount, currency, sales_owner, onboarding_owner, notes. Exclude passwords, API tokens, card/bank credentials and private keys.")
bullet(doc, "Upload must support template download, dry-run validation, duplicate detection, row-level error file, approval, import batch ID and audit history.")
bullet(doc, "For a single customer, prefer the secure form. Use file intake only for authorized bulk onboarding.")

doc.add_heading("8. Credential creation and delivery options", level=1)
table(doc, ["Option", "Method", "Decision"], [
    ("Customer chooses password", "Account invitation with one-time expiring link; MFA enrollment", "Recommended"),
    ("Team creates temporary password", "Random temporary password, forced reset on first login, delivered on separate channel", "Fallback only"),
    ("SSO", "Customer identity provider (OIDC/SAML) for enterprise", "Later/enterprise"),
    ("Magic link", "Short-lived email authentication link", "Optional after security review"),
])
doc.add_paragraph("Credentials policy: unique accounts; role-based permissions; MFA for owners/admins; invite expiry; account lock/rate limiting; password reset; session revocation; access audit; staff impersonation only with consent and logging; immediate deprovisioning on cancellation or staff exit. Integration tokens must be encrypted and masked in the UI.")

doc.add_heading("9. Current implementation status: LeadOS and AllianceOS", level=1)
doc.add_paragraph("Status is based on the repository reviewed on 24 August 2026. ‘Partial’ means code/schema exists but it does not complete a production SaaS lifecycle.")
table(doc, ["Capability", "Current status", "Repository evidence / observation"], [
    ("LeadOS operational UI", "Present", "Authenticated routes for dashboard, leads, inbox, campaigns, templates, AI brain, reports, clients, users/settings and integrations."),
    ("AllianceOS operational UI", "Present", "Authenticated routes/API for upload, prospects, planner, email/WhatsApp campaigns, replies, knowledge/prompts and analytics."),
    ("Login", "Partial", "Email/password login and 7-day JWT exist; no public signup, verification, forgot-password flow, MFA or customer invitation workflow."),
    ("Customer/team users", "Partial", "Admin can create users, but user creation does not explicitly assign the current tenant and list queries are not tenant-filtered."),
    ("Tenant schema", "Legacy/partial", "tenants, tenant_modules, subscriptions and users.tenant_id exist in the SQL backup."),
    ("Runtime tenant isolation", "Missing/critical", "Server comments state only tenant ID 1 is used; JWT omits tenant_id; many queries are global rather than tenant-scoped."),
    ("LeadOS/AllianceOS plan catalogue", "Missing", "No live product pricing/feature administration for these modules. Thedal/Mafiya plan pages are separate and out of scope."),
    ("Module entitlements", "Missing in live flow", "Legacy tenant_modules table exists, but navigation and APIs do not enforce purchased LeadOS/AllianceOS access."),
    ("Subscription lifecycle", "Missing in live flow", "Legacy subscriptions table exists; no customer checkout, activation, renewal, expiry, suspension, cancellation or proration APIs for LeadOS/AllianceOS."),
    ("Online payments", "Partial", "Lead-based Razorpay payment-link creation and captured-payment webhook exist. They record service payments, not SaaS orders/subscriptions/entitlements."),
    ("Offline payments", "Missing", "No controlled offline receipt, proof, finance approval, allocation or reconciliation workflow located."),
    ("Automated onboarding", "Partial/non-product", "Lead payment can forward to an n8n customer-journey webhook; this is not a transactional tenant/user/module provisioning flow."),
    ("Client/integration onboarding", "Partial", "LeadOS has internal client management and Meta Embedded Signup/phone registration. It does not create a subscribing customer tenant."),
    ("Invoices/tax/refunds", "Missing or not evidenced", "No complete SaaS invoice, GST, credit-note, refund and reconciliation ledger flow located."),
    ("Expiry/renewal notifications", "Missing", "No subscription reminder, grace period, dunning or entitlement restriction process located."),
    ("Audit/security controls", "Partial", "Authentication and roles exist, but SaaS-grade tenant authorization, invite audit, offline maker-checker and subscription event log are missing."),
], widths=[1.45, 1.0, 4.75])

doc.add_heading("10. Missing-process checklist before selling subscriptions", level=1)
for heading, items in [
    ("P0 — must be complete before any external multi-customer launch", [
        "Tenant isolation for every LeadOS and AllianceOS table/query/API/socket/job/webhook/export; automated cross-tenant security tests.",
        "JWT/session includes tenant identity; server derives tenant from authenticated identity and never trusts a client-supplied tenant ID.",
        "Plan, price/version, order, subscription, entitlement, payment allocation and audit-event data model with database constraints.",
        "Server-side module and feature authorization; sidebar hiding alone is not security.",
        "Secure invitation/password reset, owner/admin roles, MFA decision, rate limiting and secret encryption.",
        "Razorpay webhook signature verification, idempotency, replay safety and transactional provisioning.",
        "Privacy/terms/DPA, data retention/export/deletion, backup/restore and incident response approved for product customers.",
    ]),
    ("P1 — required for controlled commercial operations", [
        "Public pricing/quote page, registration, checkout, payment-success/pending/failure pages and onboarding status screen.",
        "Offline payment request/proof/approval, invoice/receipt, reconciliation, refund and credit-note workflow.",
        "Subscription start/end rules, timezone, renewal, grace period, suspension, cancellation, upgrade/downgrade and proration rules.",
        "Provisioning queue, owner, SLA, integration checklist, test transaction, welcome pack and training handoff.",
        "Usage meters/limits for users, leads/prospects, messages/emails, campaigns, AI calls and storage; overage behavior.",
        "Renewal and failed-payment communications; support escalation and customer-visible billing history.",
    ]),
    ("P2 — scale and optimization", [
        "Bulk onboarding, SSO, coupons/trials, partner/reseller model, multi-currency, advanced revenue analytics and self-service plan changes.",
        "Automated integration health, customer success scoring, churn controls and SLA reporting.",
    ]),
]:
    doc.add_heading(heading, level=2)
    for item in items:
        bullet(doc, "☐ " + item)

doc.add_heading("11. Proposed implementation blueprint", level=1)
table(doc, ["Phase", "Deliverables", "Exit test"], [
    ("0 — Commercial decisions", "Approve packages, price, duration, taxes, limits, setup fees, trial/refund/grace/cancellation policies", "Signed product decision sheet and terms"),
    ("1 — SaaS foundation", "Tenant scoping, migrations, plan/price/order/subscription/entitlement/audit models, RBAC, tests", "Two test tenants cannot see or change each other’s data"),
    ("2 — Sales and billing", "Pricing/quote, registration, verification, Razorpay checkout/webhooks, offline approval, invoices", "Paid and offline-approved test orders reconcile exactly once"),
    ("3 — Provisioning", "Tenant/user invite, module enablement, onboarding case/checklist, integrations, welcome notifications", "A new customer reaches active without direct database edits"),
    ("4 — Lifecycle", "Renewal, retries, reminders, grace, suspend/reactivate/cancel, upgrade/downgrade, data policy", "Date/payment scenarios pass automated acceptance tests"),
    ("5 — Launch readiness", "Security review, load/backup/restore, monitoring, support runbooks, pilot customers", "Pilot sign-off and rollback procedure approved"),
])

doc.add_heading("12. Decision sheet — complete before implementation", level=1)
table(doc, ["Decision", "Select / fill", "Owner / approval"], [
    ("Packages", "☐ LeadOS  ☐ AllianceOS  ☐ Growth Suite  ☐ Custom", "________________"),
    ("Published tiers", "☐ Starter  ☐ Growth  ☐ Scale  ☐ Enterprise", "________________"),
    ("Approved monthly prices", "LeadOS ₹_____  AllianceOS ₹_____  Bundle ₹_____", "________________"),
    ("Terms", "☐ 1 month  ☐ 3 months  ☐ 1 year  ☐ 2 years  ☐ Custom", "________________"),
    ("Discounts", "3m ____% | 1y ____% | 2y ____%", "________________"),
    ("Setup fee", "₹_____ / ☐ waived by approval", "________________"),
    ("Tax display", "☐ GST inclusive  ☐ GST extra", "________________"),
    ("Payment", "☐ Online one-time  ☐ Recurring  ☐ Bank/UPI  ☐ Cheque/cash", "________________"),
    ("Onboarding", "☐ Hybrid  ☐ Self-service  ☐ Team form  ☐ File  ☐ Dedicated", "________________"),
    ("Activation point", "☐ Payment verified  ☐ Setup complete  ☐ Contract signed + payment", "________________"),
    ("Grace / suspension", "Grace ____ days | Read-only ____ days | Delete/archive after ____ days", "________________"),
    ("Refund/cancellation", "Notice ____ days | Refund rule ____________________", "________________"),
    ("Included usage", "Users ___ | leads/prospects ___ | messages/emails ___ | AI ___ | storage ___", "________________"),
    ("Provider charges", "☐ Included to limit  ☐ Pass-through  ☐ Customer-owned accounts", "________________"),
    ("Support", "Channels __________ | Hours __________ | SLA __________", "________________"),
])

doc.add_heading("Appendix A — Suggested core records", level=1)
doc.add_paragraph("At minimum: tenants, users, invitations, products, plans, plan_prices (versioned), plan_features, orders, order_items, subscriptions, subscription_items/module_entitlements, payments, payment_allocations, invoices, offline_payment_proofs, usage_events/counters, onboarding_cases/tasks, integration_connections, consent_acceptances and audit_events. Use exact start/end timestamps rather than only text such as ‘1 Month’. Preserve the purchased price/version so later price changes do not rewrite historical contracts.")

doc.add_heading("Appendix B — Acceptance scenarios", level=1)
for item in [
    "Two tenants use identical lead/prospect names and still see only their own data, attachments, reports, sockets and background jobs.",
    "A duplicate/replayed payment webhook creates only one payment allocation and one subscription activation.",
    "An offline payment cannot activate until an authorized second person approves it.",
    "LeadOS-only customers receive 403 from AllianceOS APIs even if they manually enter the URL; the reverse also applies.",
    "Expiry at the defined timezone moves the account through grace/restricted/suspended states exactly as approved.",
    "Upgrade/downgrade, cancellation, refund and reactivation preserve a complete audit trail and invoice correctness.",
    "Invitation links expire, cannot be reused, and require password setup/MFA according to policy.",
    "No exported file, log, email, error response or analytics query leaks another tenant’s data or integration secret.",
]:
    bullet(doc, "☐ " + item)

footer = doc.sections[0].footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("LeadOS + AllianceOS subscription product decision draft — confidential")

doc.save(OUT)
print(OUT)
