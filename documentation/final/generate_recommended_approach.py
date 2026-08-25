from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("Leados-AllianceOS-Recommended-Subscription-Approach.docx")


def shade(cell, color):
    props = cell._tc.get_or_add_tcPr()
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), color)
    props.append(element)


def cell_text(cell, value, bold=False, color=None, size=9):
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(value))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        cell_text(tbl.rows[0].cells[index], header, True, "FFFFFF")
        shade(tbl.rows[0].cells[index], "1F4E78")
    for row_number, row in enumerate(rows):
        cells = tbl.add_row().cells
        for index, value in enumerate(row):
            cell_text(cells[index], value)
            if row_number % 2:
                shade(cells[index], "EAF2F8")
    doc.add_paragraph()


def bullet(doc, value, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(value)


def step(doc, title, detail):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(title + ": ")
    run.bold = True
    p.add_run(detail)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)

doc.styles["Normal"].font.name = "Aptos"
doc.styles["Normal"].font.size = Pt(10)
doc.styles["Title"].font.name = "Aptos Display"
doc.styles["Title"].font.color.rgb = RGBColor(31, 78, 120)
for name in ("Heading 1", "Heading 2"):
    doc.styles[name].font.name = "Aptos Display"
    doc.styles[name].font.color.rgb = RGBColor(31, 78, 120)

p = doc.add_paragraph(style="Title")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("LeadOS + AllianceOS\nRecommended Subscription Approach")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Simple decision and implementation guide")
r.bold = True
r.font.size = Pt(12)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run(date.today().strftime("%d %B %Y"))

notice = doc.add_table(rows=1, cols=1)
notice.style = "Table Grid"
shade(notice.cell(0, 0), "FFF2CC")
cell_text(notice.cell(0, 0), "Scope: Sell LeadOS and AllianceOS only. Mafiya OS, Thedal OS and Content OS are not included. Prices below are suggested starting points and require final management approval.", True, "7F6000", 10)

doc.add_heading("1. Recommended business model", level=1)
doc.add_paragraph("Launch as a multi-tenant subscription product using assisted automatic onboarding. Customers should be able to register, select a plan and pay online, while your team verifies integrations and completes final activation. This is safer for the first launch than completely automatic self-service because WhatsApp, Meta, email and domain setup may require manual verification.")

add_table(doc, ["Area", "Recommended choice"], [
    ("Products", "LeadOS, AllianceOS and a combined Growth Suite"),
    ("Plans", "Starter, Growth and Scale; custom Enterprise quotation"),
    ("Terms", "1 month, 3 months and 1 year; custom term when required"),
    ("Payment", "Razorpay online plus finance-approved bank/UPI transfer"),
    ("Onboarding", "Customer registration and payment followed by team-assisted integration setup"),
    ("Credentials", "Secure, expiring invitation link so the client creates their own password"),
    ("Activation", "Only after payment verification and onboarding checklist completion"),
])

doc.add_heading("2. Recommended packages and starting prices", level=1)
add_table(doc, ["Product", "Starter / month", "Growth / month", "Scale / month"], [
    ("LeadOS", "₹7,999", "₹14,999", "₹29,999"),
    ("AllianceOS", "₹9,999", "₹17,999", "₹34,999"),
    ("Growth Suite: both modules", "₹14,999", "₹27,999", "₹54,999"),
    ("Enterprise", "Custom quotation", "Custom quotation", "Custom quotation"),
])
bullet(doc, "Starter: small team with limited users, leads/prospects, campaigns and messages.")
bullet(doc, "Growth: normal business usage with more users, automation and reporting.")
bullet(doc, "Scale: higher limits, priority support and advanced integration requirements.")
bullet(doc, "Charge a one-time setup fee of approximately ₹5,000–₹25,000 when migration, integrations, training or customization is required.")
bullet(doc, "Confirm GST treatment and provider usage charges before publishing prices.")

doc.add_heading("3. Recommended subscription duration", level=1)
add_table(doc, ["Duration", "Recommended pricing", "Purpose"], [
    ("1 month", "Normal monthly price", "Low-commitment entry option"),
    ("3 months", "5% discount", "Encourages short-term commitment"),
    ("1 year", "15% discount", "Recommended best-value plan"),
    ("2 years", "Offer only by quotation initially", "Use after retention and service costs are understood"),
    ("Custom", "Approved quotation", "Enterprise usage, installments or special scope"),
])

doc.add_heading("4. Recommended payment process", level=1)
bullet(doc, "Online: use Razorpay checkout or payment link. Activate only after a verified webhook/signature.")
bullet(doc, "Offline: allow bank or UPI transfer. Client uploads proof or sends the transaction reference; finance verifies and approves it.")
bullet(doc, "Do not automatically activate from a screenshot or unverified payment reference.")
bullet(doc, "Store order, invoice, payment, subscription and entitlement records separately.")
bullet(doc, "Send invoice/receipt after approval and reminders before renewal or expiry.")

doc.add_heading("5. Recommended client onboarding flow", level=1)
for title, detail in [
    ("Choose product", "Client selects LeadOS, AllianceOS or the combined Growth Suite."),
    ("Choose plan and term", "Client selects Starter/Growth/Scale and 1 month/3 months/1 year."),
    ("Register", "Collect company, owner, billing, GST, expected usage and integration-readiness details."),
    ("Verify", "Verify the customer email and, where required, mobile number."),
    ("Accept agreements", "Record acceptance of the current Terms, Privacy Policy and data-processing terms."),
    ("Collect payment", "Use Razorpay or create an offline payment verification request."),
    ("Create account", "After payment approval, create the tenant, subscription and purchased module permissions."),
    ("Invite client", "Send a one-time expiring link so the owner creates a password; never email a permanent password."),
    ("Configure integrations", "Your team assists with Meta/WhatsApp, email, domain, branding, imports and AI knowledge."),
    ("Test and activate", "Complete test lead/message/campaign checks, then mark the account active."),
    ("Support and renew", "Send training/support information and renewal notices 30, 15, 7 and 1 day before expiry."),
]:
    step(doc, title, detail)

doc.add_heading("6. Information to collect from the client", level=1)
add_table(doc, ["Section", "Information"], [
    ("Owner", "Name, work email, mobile and designation"),
    ("Company", "Legal/display name, address, website, industry and GST details when applicable"),
    ("Subscription", "Product, tier, duration, start date, users and expected usage"),
    ("Billing", "Billing name/address/email, GST treatment and payment method"),
    ("Integrations", "WhatsApp/Meta readiness, business number, email/domain and import requirements"),
    ("Consent", "Terms, privacy and data-processing acceptance with version and timestamp"),
])
doc.add_paragraph("For bulk client creation, use a validated CSV/XLSX template. Never include passwords, API tokens, private keys or card/bank credentials in the file.")

doc.add_heading("7. What must be implemented before launch", level=1)
for item in [
    "Separate every client’s data using secure tenant-level filtering.",
    "Add LeadOS and AllianceOS plan and price management.",
    "Enable only the modules and features purchased by each customer.",
    "Build registration, email verification, pricing, checkout and onboarding-status pages.",
    "Connect verified payment to automatic subscription and entitlement creation.",
    "Add secure invitation, password reset, role permissions and preferably MFA for administrators.",
    "Implement offline payment approval, invoice/GST, reconciliation, refund and audit processes.",
    "Implement renewal, failed-payment retry, grace period, expiry, suspension and reactivation.",
    "Create usage limits for users, leads/prospects, messages/emails, campaigns, AI and storage.",
    "Test that one customer can never access another customer’s records, files, reports or integrations.",
]:
    bullet(doc, "☐ " + item)

doc.add_heading("8. Recommended launch phases", level=1)
add_table(doc, ["Phase", "Action", "Result"], [
    ("Phase 1", "Approve prices, plans, limits, tax, refund, grace period and support policy", "Commercial rules are fixed"),
    ("Phase 2", "Implement tenant isolation and module permissions", "External customer access becomes safe"),
    ("Phase 3", "Build registration, payment, subscription and account invitation", "Customer can purchase and receive an account"),
    ("Phase 4", "Build assisted integration checklist and activation process", "Team can onboard consistently"),
    ("Phase 5", "Add renewal, suspension, billing history, usage tracking and reporting", "Complete subscription lifecycle"),
    ("Phase 6", "Pilot with 3–5 customers before public launch", "Problems are found with controlled risk"),
])

doc.add_heading("9. Final recommendation", level=1)
p = doc.add_paragraph()
r = p.add_run("Start with assisted onboarding and three subscription terms: 1 month, 3 months and 1 year. ")
r.bold = True
p.add_run("Offer separate LeadOS and AllianceOS plans plus a discounted combined bundle. Accept Razorpay and finance-approved bank/UPI payments. Pilot with 3–5 clients. Move to full self-service only after tenant isolation, automatic entitlement provisioning, renewal and suspension processes have passed security and acceptance testing.")

doc.add_heading("10. Management approval", level=1)
add_table(doc, ["Decision", "Selection"], [
    ("Approved products", "☐ LeadOS  ☐ AllianceOS  ☐ Growth Suite"),
    ("Approved terms", "☐ 1 month  ☐ 3 months  ☐ 1 year  ☐ Custom"),
    ("Approved monthly prices", "LeadOS ₹_____ | AllianceOS ₹_____ | Bundle ₹_____"),
    ("Setup fee", "₹_____ or approval rule: ____________________"),
    ("Payment modes", "☐ Razorpay  ☐ Bank/UPI  ☐ Other: __________"),
    ("Grace period", "_____ days"),
    ("Approved by", "Name/signature/date: ______________________________"),
])

footer = doc.sections[0].footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("LeadOS + AllianceOS recommended subscription approach — confidential")

doc.save(OUT)
print(OUT)
